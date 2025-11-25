from flask import Flask, render_template, request
import re
import os
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document


app = Flask(__name__)

# Upload configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'.txt', '.pdf', '.docx'}
MAX_FILE_SIZE = 16 * 1024 * 1024  # 16 MB

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

# Create upload folder if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# PII Risk Patterns Dictionary
PII_RISK_PATTERNS = {
    "Email Address": {
        "pattern": r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,6}",
        "risk": "High"
    },
    "Credit Card (VISA/Mastercard)": {
        "pattern": r"(?:4\d{3}|5[1-5]\d{2})[ -]?\d{4}[ -]?\d{4}[ -]?\d{4}",
        "risk": "High"
    },
    "US Social Security Number (SSN)": {
        "pattern": r"\b\d{3}[ -]?\d{2}[ -]?\d{4}\b",
        "risk": "High"
    },
    "Passport Number": {
        "pattern": r"\b[A-Z]{1,2}\d{6,9}\b",
        "risk": "High"
    },
    "IBAN (Bank Account)": {
        "pattern": r"\b[A-Z]{2}\d{2}[ ]?[A-Z0-9]{4}[ ]?\d{4}[ ]?\d{4}[ ]?\d{4}[ ]?\d{0,4}\b",
        "risk": "High"
    },
    "Date of Birth": {
        "pattern": r"\b(?:DOB|Date of Birth|Birth Date)[:\s]+(?:\d{1,2}[-/]\d{1,2}[-/]\d{2,4})",
        "risk": "High"
    },
    "MAC Address": {
        "pattern": r"\b(?:[0-9A-Fa-f]{2}[:-]){5}[0-9A-Fa-f]{2}\b",
        "risk": "Medium"
    },
    "IPv4 Address": {
        "pattern": r"\b(?:\d{1,3}\.){3}\d{1,3}\b",
        "risk": "Medium"
    },
    "Phone Number (US/Simple Intl)": {
        "pattern": r"\b(?:\d{3}[-.\s]?){2}\d{4}\b",
        "risk": "Medium"
    }
}

# Compliance Keyword Patterns Dictionary
COMPLIANCE_KEYWORDS = {
    "Regulatory Frameworks": [
        "GDPR",
        "CCPA"
    ],
    "GDPR Compliance": [
        "Right to be Forgotten",
        "Right to Rectification",
        "Data Subject Request",
        "DSAR",
        "Lawful Basis",
        "Legitimate Interest",
        "Data Protection Officer",
        "DPO",
        "Controller",
        "Processor",
        "Joint Controller",
        "Privacy by Design",
        "Privacy by Default",
        "Data Protection Impact Assessment",
        "DPIA",
        "Supervisory Authority",
        "Consent Withdrawal",
        "Profiling",
        "Automated Decision-Making"
    ],
    "CCPA Compliance": [
        "Right to Opt-Out",
        "Do Not Sell My Personal Information",
        "California Consumer Privacy Act",
        "Sale of Personal Information",
        "Consumer Request",
        "Authorized Agent",
        "Financial Incentive",
        "Notice at Collection",
        "Categories of Personal Information"
    ],
    "Data Rights (General)": [
        "Right to Access",
        "Data Portability"
    ],
    "Policy/Consent": [
        "Cookie Policy",
        "Privacy Notice",
        "Explicit Consent",
        "Data Processing Agreement",
        "DPA"
    ],
    "Security/Transfers": [
        "Data Breach Notification",
        "Data Transfer",
        "Third Parties",
        "Security Measures",
        "Encryption",
        "Anonymization"
    ],
    "Compliance Timeframes": [
        "within 72 hours",
        "72 hours",
        "30 days",
        "within 30 days"
    ]
}


def allowed_file(filename):
    """Check if file extension is allowed."""
    return any(filename.lower().endswith(ext) for ext in ALLOWED_EXTENSIONS)


def extract_text_from_file(filepath, filename):
    """Extract text from uploaded file based on extension."""
    extension = os.path.splitext(filename.lower())[1]
    
    try:
        if extension in ['.txt']:
            # Plain text files
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif extension == '.pdf':
            # PDF files
            text = ""
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        
        elif extension == '.docx':
            # Word documents
            doc = Document(filepath)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        
        else:
            return None  # Unsupported format
            
    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
        return None


@app.route("/")
def index():
    """Serve the main input form."""
    return render_template("index.html")


@app.route("/scan", methods=["POST"])
def scan():
    """Process the text and perform PII and compliance scanning."""
    text = ""
    error_message = None
    
    # Check if file was uploaded
    if 'file' in request.files and request.files['file'].filename != '':
        file = request.files['file']
        
        # Validate file
        if not allowed_file(file.filename):
            error_message = "Invalid file type. Only .txt, .pdf, and .docx files are allowed."
        else:
            # Save file temporarily
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            
            try:
                file.save(filepath)
                
                # Extract text from file
                text = extract_text_from_file(filepath, filename)
                
                # Clean up temporary file
                os.remove(filepath)
                
                if text is None or text.strip() == "":
                    error_message = "Could not extract text from file. Please ensure it's not empty or corrupted."
                    
            except Exception as e:
                error_message = f"Error processing file: {str(e)}"
                # Clean up file if it exists
                if os.path.exists(filepath):
                    os.remove(filepath)
    
    else:
        # Get text from textarea
        text = request.form.get("text", "")
    
    # Check if we have text to scan
    if not text or text.strip() == "":
        if not error_message:
            error_message = "Please provide text to scan or upload a file."
        return render_template("index.html", error=error_message)
    
    # If there was an error during file processing, show it
    if error_message:
        return render_template("index.html", error=error_message)
    
    # Split text into lines for line number tracking
    lines = text.split("\n")
    
    # PII Risk Findings
    risk_findings = []
    
    # Scan for PII patterns
    for pii_type, pii_info in PII_RISK_PATTERNS.items():
        pattern = pii_info["pattern"]
        risk_level = pii_info["risk"]
        
        # Search each line for matches
        for line_num, line in enumerate(lines, start=1):
            matches = re.finditer(pattern, line)
            for match in matches:
                # Extract snippet with context (up to 50 chars around match)
                start = max(0, match.start() - 10)
                end = min(len(line), match.end() + 10)
                snippet = line[start:end]
                
                risk_findings.append({
                    "type": pii_type,
                    "risk_level": risk_level,
                    "line_number": line_num,
                    "snippet": snippet,
                    "matched_text": match.group()
                })
    
    # Calculate Risk Severity and Critical Patterns
    critical_alerts = []
    pii_type_counts = {}
    
    # Count each PII type
    for finding in risk_findings:
        pii_type = finding["type"]
        pii_type_counts[pii_type] = pii_type_counts.get(pii_type, 0) + 1
    
    # Detect critical patterns
    has_ssn = "US Social Security Number (SSN)" in pii_type_counts
    has_cc = "Credit Card (VISA/Mastercard)" in pii_type_counts
    has_email = "Email Address" in pii_type_counts
    has_passport = "Passport Number" in pii_type_counts
    has_iban = "IBAN (Bank Account)" in pii_type_counts
    has_dob = "Date of Birth" in pii_type_counts
    
    # Identity Theft Risk (SSN + Credit Card)
    if has_ssn and has_cc:
        critical_alerts.append({
            "icon": "fa-skull-crossbones",
            "title": "Identity Theft Risk",
            "description": f"Found SSN ({pii_type_counts.get('US Social Security Number (SSN)', 0)}) and Credit Card ({pii_type_counts.get('Credit Card (VISA/Mastercard)', 0)}) - Complete identity profile exposed",
            "severity": "critical"
        })
    
    # Financial Fraud Risk (Credit Card + Email or IBAN + Email)
    if (has_cc or has_iban) and has_email:
        financial_items = []
        if has_cc:
            financial_items.append(f"Credit Cards ({pii_type_counts.get('Credit Card (VISA/Mastercard)', 0)})")
        if has_iban:
            financial_items.append(f"Bank Accounts ({pii_type_counts.get('IBAN (Bank Account)', 0)})")
        critical_alerts.append({
            "icon": "fa-credit-card",
            "title": "Financial Fraud Risk",
            "description": f"Found {', '.join(financial_items)} with Email ({pii_type_counts.get('Email Address', 0)}) - Financial fraud potential",
            "severity": "critical"
        })
    
    # Full Identity Profile Risk (SSN/Passport + DOB + Email)
    if (has_ssn or has_passport) and has_dob and has_email:
        critical_alerts.append({
            "icon": "fa-user-secret",
            "title": "Full Identity Profile",
            "description": "Found combination of SSN/Passport, Date of Birth, and Email - Complete identity exposure",
            "severity": "critical"
        })
    
    # Mass Data Exposure
    total_pii = len(risk_findings)
    if total_pii >= 50:
        critical_alerts.append({
            "icon": "fa-database",
            "title": "Mass Data Exposure",
            "description": f"Found {total_pii} PII items - Large-scale data exposure risk",
            "severity": "critical"
        })
    elif total_pii >= 20:
        critical_alerts.append({
            "icon": "fa-exclamation-triangle",
            "title": "High Volume Data Exposure",
            "description": f"Found {total_pii} PII items - Significant data exposure",
            "severity": "warning"
        })
    
    # Assign severity status to each finding based on critical patterns
    for finding in risk_findings:
        finding_type = finding["type"]
        
        # Critical if part of identity theft combo
        if finding_type in ["US Social Security Number (SSN)", "Credit Card (VISA/Mastercard)"] and has_ssn and has_cc:
            finding["severity"] = "Critical"
            finding["severity_icon"] = "fa-circle-xmark"
            finding["severity_class"] = "critical"
        # Critical if financial data with contact info
        elif finding_type in ["Credit Card (VISA/Mastercard)", "IBAN (Bank Account)"] and has_email:
            finding["severity"] = "Critical"
            finding["severity_icon"] = "fa-circle-xmark"
            finding["severity_class"] = "critical"
        # High risk by default (already high risk items)
        elif finding["risk_level"] == "High":
            finding["severity"] = "High"
            finding["severity_icon"] = "fa-circle-exclamation"
            finding["severity_class"] = "high"
        # Medium risk
        else:
            finding["severity"] = "Medium"
            finding["severity_icon"] = "fa-circle-info"
            finding["severity_class"] = "medium"
    
    # Count severity levels
    severity_counts = {
        "critical": sum(1 for f in risk_findings if f.get("severity") == "Critical"),
        "high": sum(1 for f in risk_findings if f.get("severity") == "High"),
        "medium": sum(1 for f in risk_findings if f.get("severity") == "Medium")
    }
    
    # Compliance Keyword Status
    compliance_status = []
    
    # Check for compliance keywords (case-insensitive)
    text_lower = text.lower()
    
    for category, keywords in COMPLIANCE_KEYWORDS.items():
        found_keywords = []
        missing_keywords = []
        
        for keyword in keywords:
            if keyword.lower() in text_lower:
                found_keywords.append(keyword)
            else:
                missing_keywords.append(keyword)
        
        compliance_status.append({
            "category": category,
            "found": found_keywords,
            "missing": missing_keywords,
            "has_any": len(found_keywords) > 0
        })
    
    # Render results template with findings
    return render_template("report.html", 
                         risk_findings=risk_findings, 
                         compliance_status=compliance_status,
                         total_risks=len(risk_findings),
                         critical_alerts=critical_alerts,
                         severity_counts=severity_counts,
                         pii_type_counts=pii_type_counts)


@app.errorhandler(413)
def request_entity_too_large(error):
    """Handle file too large error."""
    return render_template("index.html", 
                         error="File is too large! Maximum file size is 16 MB. Please upload a smaller file."), 413


if __name__ == "__main__":
    app.run(debug=True)
